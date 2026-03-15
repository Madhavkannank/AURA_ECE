import os
import re
import tempfile
from collections import Counter
from pathlib import Path
from typing import Any
from uuid import uuid4

from fastapi import UploadFile

from ..config import get_settings
from .groq_client import get_groq_service

try:
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = None

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None

try:
    import pytesseract
except ImportError:  # pragma: no cover
    pytesseract = None

try:
    from rapidocr_onnxruntime import RapidOCR
except ImportError:  # pragma: no cover
    RapidOCR = None

try:
    import pikepdf
except ImportError:  # pragma: no cover
    pikepdf = None


class NotesService:
    def __init__(self) -> None:
        self.settings = get_settings()
        self.groq = get_groq_service()
        self.storage_dir = Path(self.settings.notes_storage_path).resolve()
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        self._rapid_ocr = RapidOCR() if RapidOCR else None

    def _safe_name(self, name: str) -> str:
        base = os.path.basename(name or "note")
        return re.sub(r"[^A-Za-z0-9._-]", "_", base)

    def detect_file_kind(self, file_name: str | None, content_type: str | None) -> str:
        suffix = Path(file_name or "").suffix.lower()
        ctype = (content_type or "").lower()

        video_exts = {".mp4", ".mov", ".avi", ".mkv", ".webm", ".m4v"}
        audio_exts = {".wav", ".mp3", ".m4a", ".aac", ".ogg", ".flac"}
        image_exts = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"}
        doc_exts = {".pdf", ".docx", ".txt", ".md", ".csv", ".log"}

        if ctype.startswith("video/") or suffix in video_exts:
            return "video"
        if ctype.startswith("audio/") or suffix in audio_exts:
            return "audio"
        if ctype.startswith("image/") or suffix in image_exts:
            return "image"
        if ctype in {"application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/plain"} or suffix in doc_exts:
            return "document"
        return "other"

    async def save_upload(self, file: UploadFile, file_kind: str = "document") -> Path:
        original = self._safe_name(file.filename or "note")
        target_dir = self.storage_dir / file_kind
        target_dir.mkdir(parents=True, exist_ok=True)
        stored = target_dir / f"{uuid4().hex[:10]}_{original}"
        data = await file.read()
        with open(stored, "wb") as handle:
            handle.write(data)
        return stored

    def extract_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        text = ""

        if suffix == ".pdf":
            text = self._extract_pdf(file_path)
        elif suffix == ".docx":
            text = self._extract_docx(file_path)
        elif suffix in {".txt", ".md", ".csv", ".log"}:
            text = self._extract_text_file(file_path)
        elif suffix in {".png", ".jpg", ".jpeg", ".bmp", ".tiff"}:
            text = self._extract_image_ocr(file_path)

        clean = re.sub(r"\s+", " ", text or "").strip()
        if not clean:
            return ""
        return clean[: max(1000, int(self.settings.notes_max_chars))]

    def _extract_pdf(self, file_path: Path) -> str:
        if not pdfplumber:
            return ""
        content: list[str] = []
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    content.append(page_text)
                    continue

                # Fallback for scanned/image-only pages.
                content.append(self._extract_pdf_page_ocr(page))
        return "\n".join(content)

    def _extract_pdf_page_ocr(self, page: Any) -> str:
        if not Image:
            return ""

        try:
            rendered = page.to_image(resolution=200).original
        except Exception:
            return ""

        if pytesseract:
            try:
                return pytesseract.image_to_string(rendered)
            except Exception:
                pass

        if self._rapid_ocr:
            try:
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    tmp_path = Path(tmp.name)
                try:
                    rendered.save(str(tmp_path), format="PNG")
                    result, _ = self._rapid_ocr(str(tmp_path))
                    if not result:
                        return ""
                    lines = [str(item[1]) for item in result if len(item) > 1]
                    return "\n".join(lines)
                finally:
                    try:
                        tmp_path.unlink(missing_ok=True)
                    except Exception:
                        pass
            except Exception:
                return ""

        return ""

    def _extract_docx(self, file_path: Path) -> str:
        if not Document:
            return ""
        doc = Document(str(file_path))
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    def _extract_text_file(self, file_path: Path) -> str:
        try:
            return file_path.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            return ""

    def _extract_image_ocr(self, file_path: Path) -> str:
        # Prefer Tesseract when available, fall back to ONNX-based OCR.
        if pytesseract and Image:
            try:
                img = Image.open(str(file_path))
                return pytesseract.image_to_string(img)
            except Exception:
                pass

        if self._rapid_ocr:
            try:
                result, _ = self._rapid_ocr(str(file_path))
                if not result:
                    return ""
                lines = [str(item[1]) for item in result if len(item) > 1]
                return "\n".join(lines)
            except Exception:
                return ""

        return ""

    def _chunk_text(self, text: str) -> list[str]:
        chunk_size = max(1000, int(self.settings.notes_chunk_chars))
        if len(text) <= chunk_size:
            return [text]
        return [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)]

    def analyze_text(self, text: str) -> dict[str, Any]:
        if not text:
            return {
                "keywords": ["unclassified"],
                "category": "Unclassified",
                "summary": "No extractable text was found in the file.",
            }

        chunks = self._chunk_text(text)
        all_keywords: list[str] = []
        categories: list[str] = []

        for chunk in chunks[:6]:
            prompt = (
                "Analyze this educational note/document and return JSON with: "
                "category (single short label), keywords (5-10 concise terms), summary (1-2 sentences). "
                "Focus on school note relevance for student/teacher records. "
                f"Document text: {chunk}"
            )
            data = self.groq.chat_json(prompt, self.groq.settings.groq_light_model)
            if not isinstance(data, dict):
                continue

            category = str(data.get("category", "")).strip()
            if category:
                categories.append(category)

            keywords = data.get("keywords", [])
            if isinstance(keywords, list):
                all_keywords.extend([str(k).strip() for k in keywords if str(k).strip()])
            elif isinstance(keywords, str):
                all_keywords.extend([k.strip() for k in keywords.split(",") if k.strip()])

        if not all_keywords:
            fallback = [w.lower() for w in re.findall(r"[A-Za-z]{4,}", text)[:30]]
            all_keywords = fallback or ["education", "notes"]

        top_keywords = [k for k, _ in Counter([k.lower() for k in all_keywords]).most_common(10)]
        category = Counter(categories).most_common(1)[0][0] if categories else "General Note"

        summary_prompt = (
            "Create a concise summary (max 2 sentences) of this educational note for searchable metadata. "
            "Return plain text only. "
            f"Text: {text[:6000]}"
        )
        summary = self.groq.chat_text(summary_prompt, self.groq.settings.groq_reasoning_model)
        if not summary:
            summary = f"{category} note with key topics: {', '.join(top_keywords[:5])}."

        return {
            "keywords": top_keywords,
            "category": category,
            "summary": summary.strip(),
        }

    def embed_metadata(self, file_path: Path, keywords: list[str], category: str) -> bool:
        suffix = file_path.suffix.lower()
        keyword_str = ", ".join(keywords)

        if suffix == ".pdf" and pikepdf:
            try:
                with pikepdf.open(str(file_path)) as pdf:
                    pdf.docinfo["/Keywords"] = keyword_str
                    pdf.docinfo["/Subject"] = category
                    pdf.save(str(file_path))
                return True
            except Exception:
                return False

        if suffix == ".docx" and Document:
            try:
                doc = Document(str(file_path))
                doc.core_properties.keywords = keyword_str
                doc.core_properties.subject = category
                doc.save(str(file_path))
                return True
            except Exception:
                return False

        return False


notes_service = NotesService()
